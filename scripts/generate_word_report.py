"""
Mami AI Enterprise Evaluation Report - Word Document Generator
Generates a professional Word document from the analysis.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import datetime

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_with_style(doc, text, level=1):
    """Add a styled heading."""
    heading = doc.add_heading(text, level=level)
    return heading

def add_table_with_style(doc, headers, rows, col_widths=None):
    """Add a styled table."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], "2E86AB")
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    # Data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = str(cell_text)
    
    doc.add_paragraph()
    return table

def create_enterprise_report():
    """Create the enterprise evaluation Word document."""
    doc = Document()
    
    # Set document properties
    core_props = doc.core_properties
    core_props.author = "Antigravity AI Assistant"
    core_props.title = "Mami AI v4.2 - Enterprise Değerlendirme Raporu"
    core_props.subject = "10 Premium Müşteri için Kurumsal AI Asistanı Değerlendirmesi"
    
    # =========================================================================
    # TITLE PAGE
    # =========================================================================
    title = doc.add_heading("Mami AI v4.2", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph("Enterprise Değerlendirme Raporu")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(24)
    subtitle.runs[0].font.color.rgb = RGBColor(46, 134, 171)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("Hazırlık Tarihi: ").bold = True
    info.add_run(f"{datetime.datetime.now().strftime('%d %B %Y')}\n")
    info.add_run("Versiyon: ").bold = True
    info.add_run("4.2.0\n")
    info.add_run("Hedef: ").bold = True
    info.add_run("10 Premium Müşteriye Kurumsal AI Hizmeti")
    
    doc.add_page_break()
    
    # =========================================================================
    # TABLE OF CONTENTS
    # =========================================================================
    add_heading_with_style(doc, "İçindekiler", 1)
    
    toc_items = [
        "1. Yönetici Özeti",
        "2. Mevcut Sistem Mimarisi",
        "3. Özellik Envanteri",
        "4. Backend Analizi",
        "5. Frontend Analizi",
        "6. Veritabanı ve Depolama",
        "7. Enterprise Gap Analizi",
        "8. Evrim Yol Haritası",
        "9. Teknik Borç ve Riskler",
        "10. Öneriler ve Sonuç"
    ]
    
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_page_break()
    
    # =========================================================================
    # 1. YÖNETİCİ ÖZETİ
    # =========================================================================
    add_heading_with_style(doc, "1. Yönetici Özeti", 1)
    
    add_heading_with_style(doc, "Proje Tanımı", 2)
    doc.add_paragraph(
        "Mami AI v4.2, Türkçe odaklı, çok modelli, hafızalı ve görsel üretim yetkinlikli "
        "bir yapay zeka asistan platformudur. FastAPI backend ve React/Vite frontend "
        "mimarisi üzerine kurulmuştur."
    )
    
    add_heading_with_style(doc, "Güçlü Yönler", 2)
    add_table_with_style(doc, 
        ["Alan", "Düzey", "Detay"],
        [
            ["LLM Entegrasyonu", "⭐⭐⭐⭐⭐", "Multi-model (Groq, Ollama), akıllı routing"],
            ["Hafıza Sistemi", "⭐⭐⭐⭐", "4-katmanlı mimari (Working/Profile/Semantic/Archive)"],
            ["RAG Sistemi", "⭐⭐⭐⭐", "Page-aware, hybrid search, multilingual"],
            ["Görsel Üretim", "⭐⭐⭐⭐", "Flux/Forge entegrasyonu, NSFW routing"],
            ["Frontend", "⭐⭐⭐⭐", "Modern React, streaming, responsive, PWA"]
        ]
    )
    
    add_heading_with_style(doc, "Eksik Alanlar (Enterprise İçin)", 2)
    add_table_with_style(doc,
        ["Alan", "Mevcut", "Hedef"],
        [
            ["Multi-Tenant Mimari", "❌ Yok", "10 müşteri izolasyonu"],
            ["Ölçeklenebilirlik", "Tek sunucu", "Kubernetes/Load balancing"],
            ["Güvenlik", "Temel", "SOC 2 / ISO 27001 uyumu"],
            ["Monitoring", "%30", "Tam observability stack"],
            ["SLA Garantisi", "Yok", "%99.9 uptime"]
        ]
    )
    
    doc.add_page_break()
    
    # =========================================================================
    # 2. MEVCUT SİSTEM MİMARİSİ
    # =========================================================================
    add_heading_with_style(doc, "2. Mevcut Sistem Mimarisi", 1)
    
    doc.add_paragraph(
        "Sistem, modern mikroservis benzeri bir mimari üzerine kurulmuştur. "
        "Frontend (React+Vite) → API Gateway (FastAPI) → Orchestrator → LLM/Tools/Memory "
        "katmanları şeklinde organize edilmiştir."
    )
    
    add_heading_with_style(doc, "Teknoloji Stack", 2)
    add_table_with_style(doc,
        ["Katman", "Teknoloji", "Detay"],
        [
            ["Backend Framework", "FastAPI", "0.104+"],
            ["Runtime", "Python", "3.11+"],
            ["Frontend Framework", "React + Vite", "18.x"],
            ["State Management", "Zustand", "Modern, lightweight"],
            ["Primary LLM", "Groq Cloud", "llama-3.3-70b-versatile"],
            ["Local LLM", "Ollama", "josiefied-qwen3-8b"],
            ["Vector DB", "ChromaDB", "Multilingual embeddings"],
            ["Cache/Session", "Redis", "Working memory + RAG cache"],
            ["Relational DB", "SQLite", "SQLModel ORM"],
            ["Image Gen", "Stable Diffusion", "Forge/Flux"]
        ]
    )
    
    doc.add_page_break()
    
    # =========================================================================
    # 3. ÖZELLİK ENVANTERİ
    # =========================================================================
    add_heading_with_style(doc, "3. Özellik Envanteri", 1)
    
    add_heading_with_style(doc, "3.1 Yapay Zeka Yetenekleri", 2)
    add_table_with_style(doc,
        ["Özellik", "Durum", "Detay"],
        [
            ["Multi-LLM Desteği", "✅", "Groq (4 API key rotation), Ollama"],
            ["Akıllı Model Routing", "✅", "Capability-based selection"],
            ["Intent Classification", "✅", "CHAT/IMAGE/INTERNET/LOCAL ayrımı"],
            ["Streaming Responses", "✅", "Real-time SSE"],
            ["Specialist-Stylist Pipeline", "✅", "İki aşamalı yanıt kalitesi"],
            ["Multi-Intent (tasks[])", "✅", "Tek mesajda birden fazla görev"],
            ["Proactive Suggestions", "🔶", "Kısmen implemente"]
        ]
    )
    
    add_heading_with_style(doc, "3.2 Hafıza Sistemleri", 2)
    add_table_with_style(doc,
        ["Katman", "Teknoloji", "TTL", "İşlev"],
        [
            ["Layer 1: Working Memory", "Redis", "48 saat", "Son 10 mesaj + session summary"],
            ["Layer 2: User Profile", "SQLite + LLM", "Kalıcı", "Structured facts"],
            ["Layer 3: Semantic Memory", "ChromaDB", "Decay", "Uzun vadeli, duplicate detection"],
            ["Layer 4: Conversation Archive", "SQLite", "Kalıcı", "Tüm sohbet özetleri"]
        ]
    )
    
    add_heading_with_style(doc, "3.3 RAG (Retrieval-Augmented Generation)", 2)
    add_table_with_style(doc,
        ["Özellik", "Durum", "Detay"],
        [
            ["PDF/TXT Ingestion", "✅", "Page-aware, semantic chunking"],
            ["Multilingual Embeddings", "✅", "paraphrase-multilingual-MiniLM-L12-v2"],
            ["Hybrid Search", "✅", "Vector (%70) + BM25 (%30)"],
            ["Retrieval Grading", "✅", "Score > 0.7 filter"],
            ["Neighbor Expansion", "✅", "+/- 1 chunk context"],
            ["Intelligent Gate", "✅", "LLM/Web/RAG otomatik seçimi"]
        ]
    )
    
    add_heading_with_style(doc, "3.4 İnternet Araması", 2)
    add_table_with_style(doc,
        ["Provider", "Durum", "Kullanım"],
        [
            ["DuckDuckGo", "✅", "Primary (ücretsiz)"],
            ["Bing Search API", "✅", "Backup"],
            ["Serper (Google)", "✅", "Backup"]
        ]
    )
    
    add_heading_with_style(doc, "3.5 Görsel Üretim", 2)
    add_table_with_style(doc,
        ["Özellik", "Durum", "Detay"],
        [
            ["Flux/Forge Entegrasyonu", "✅", "Stable Diffusion tabanlı"],
            ["NSFW Routing", "✅", "Checkpoint seçimi"],
            ["Async Job Queue", "✅", "UUID tabanlı"],
            ["WebSocket Progress", "✅", "Real-time bildirim"],
            ["Circuit Breaker", "✅", "GPU yoğunluk koruması"]
        ]
    )
    
    add_heading_with_style(doc, "3.6 Güvenlik ve Yetkilendirme", 2)
    add_table_with_style(doc,
        ["Özellik", "Durum", "Detay"],
        [
            ["JWT Authentication", "✅", "Session-based"],
            ["Davet Kodu Sistemi", "✅", "Kontrollü erişim"],
            ["3 Seviyeli Sansür", "✅", "Unrestricted/Normal/Strict"],
            ["NSFW Detection", "✅", "Pattern-based"],
            ["Llama Guard", "🔶", "Kısmen implemente"],
            ["Prompt Injection Koruması", "🔶", "Tool-Hijack Policy var"]
        ]
    )
    
    doc.add_page_break()
    
    # =========================================================================
    # 4. BACKEND ANALİZİ
    # =========================================================================
    add_heading_with_style(doc, "4. Backend Analizi", 1)
    
    add_heading_with_style(doc, "Dizin Yapısı", 2)
    structure = doc.add_paragraph()
    structure.add_run("mami_ai_v4/\n").bold = True
    structure.add_run("""├── main.py                # Giriş noktası
├── app/                   # Ana uygulama (~15KB+ kod)
│   ├── main.py            # FastAPI app (13KB)
│   ├── config.py          # Pydantic settings (10KB)
│   ├── api/               # HTTP endpoints (13 dosya)
│   ├── auth/              # Kimlik doğrulama (7 dosya)
│   ├── chat/              # Sohbet işleme (11 dosya)
│   ├── core/              # Altyapı (22 dosya)
│   ├── image/             # Görsel üretim (9 dosya)
│   ├── memory/            # Hafıza sistemleri (12 dosya)
│   ├── orchestrator_v42/  # Ana orchestrator (70 dosya)
│   ├── plugins/           # Plugin sistemi (24 dosya)
│   └── services/          # Yardımcı servisler (16 dosya)
├── core_v2/               # Clean architecture (25 dosya)
└── tests/                 # Test suite (42 dosya)
""")
    
    add_heading_with_style(doc, "Kritik Modüller", 2)
    add_table_with_style(doc,
        ["Modül", "Satır", "İşlev"],
        [
            ["Orchestrator v4.2 (gateway.py)", "~1,344", "Ana işlem merkezi"],
            ["Chat Processor (processor.py)", "~952", "Legacy sohbet işlemcisi"],
            ["Smart Router (smart_router.py)", "~944", "Model ve tool yönlendirme"],
            ["Working Memory (working_memory.py)", "~547", "Redis tabanlı session cache"],
            ["RAG v2 (rag_v2.py)", "~1,055", "Page-aware document retrieval"]
        ]
    )
    
    doc.add_page_break()
    
    # =========================================================================
    # 5. FRONTEND ANALİZİ
    # =========================================================================
    add_heading_with_style(doc, "5. Frontend Analizi", 1)
    
    add_heading_with_style(doc, "Kritik Bileşenler", 2)
    add_table_with_style(doc,
        ["Bileşen", "Boyut", "İşlev"],
        [
            ["ChatInput.tsx", "31KB", "Ana mesaj girişi"],
            ["MessageBubble.tsx", "28KB", "Mesaj render"],
            ["DesignLabPage.tsx", "45KB", "Tasarım laboratuvarı"],
            ["ImageProgressCard.tsx", "18KB", "Görsel üretim progress"],
            ["OrchDebugPanel.tsx", "16KB", "Orchestrator debug"],
            ["MermaidViewer.tsx", "16KB", "Diagram görüntüleyici"]
        ]
    )
    
    add_heading_with_style(doc, "State Management (Zustand)", 2)
    add_table_with_style(doc,
        ["Store", "Boyut", "İşlev"],
        [
            ["chatStore.ts", "7.3KB", "Sohbet state"],
            ["settingsStore.ts", "7.9KB", "Kullanıcı ayarları"],
            ["imageJobsStore.ts", "7.5KB", "Görsel işleri"],
            ["userStore.ts", "3.7KB", "Kullanıcı bilgisi"],
            ["themeStore.ts", "2.5KB", "Tema yönetimi"]
        ]
    )
    
    add_heading_with_style(doc, "Frontend Özellikleri", 2)
    add_table_with_style(doc,
        ["Özellik", "Durum"],
        [
            ["Responsive Layout", "✅"],
            ["Dark Mode", "✅"],
            ["Streaming Yanıt", "✅"],
            ["Code Syntax Highlighting", "✅"],
            ["Mermaid Diagram Render", "✅"],
            ["Image Gallery/Lightbox", "✅"],
            ["PWA Desteği", "✅"],
            ["Command Palette (/)", "✅"],
            ["Keyboard Shortcuts", "✅"]
        ]
    )
    
    doc.add_page_break()
    
    # =========================================================================
    # 6. VERİTABANI VE DEPOLAMA
    # =========================================================================
    add_heading_with_style(doc, "6. Veritabanı ve Depolama", 1)
    
    add_heading_with_style(doc, "SQLite Tabloları", 2)
    add_table_with_style(doc,
        ["Tablo", "İşlev"],
        [
            ["User", "Kullanıcı hesapları"],
            ["Session", "Aktif oturumlar"],
            ["Conversation", "Sohbet başlıkları"],
            ["Message", "Mesaj içeriği + metadata"],
            ["Memory", "Hafıza kayıtları (meta)"],
            ["RAGDocument", "Yüklenen dokümanlar"],
            ["Feedback", "Kullanıcı geri bildirimleri"],
            ["AIIdentityConfig", "Persona yapılandırması"]
        ]
    )
    
    add_heading_with_style(doc, "Redis Key Patterns", 2)
    add_table_with_style(doc,
        ["Pattern", "TTL", "İşlev"],
        [
            ["wm:{user_id}:msgs", "48h", "Son mesajlar"],
            ["wm:{user_id}:summary", "48h", "Session özeti"],
            ["wm:{user_id}:rag:{hash}", "1h", "RAG cache"],
            ["wm:{user_id}:facts", "48h", "Anlık facts"]
        ]
    )
    
    doc.add_page_break()
    
    # =========================================================================
    # 7. ENTERPRISE GAP ANALİZİ
    # =========================================================================
    add_heading_with_style(doc, "7. Enterprise Gap Analizi", 1)
    
    add_heading_with_style(doc, "7.1 Kritik Eksiklikler (Must-Have)", 2)
    add_table_with_style(doc,
        ["Gap", "Mevcut", "Gerekli", "Öncelik"],
        [
            ["Multi-Tenant Mimari", "Tek instance", "Tenant izolasyonu", "🔴 P0"],
            ["Horizontal Scaling", "Tek sunucu", "K8s + Load Balancer", "🔴 P0"],
            ["Veritabanı", "SQLite", "PostgreSQL", "🔴 P0"],
            ["Güvenlik Sertifikasyonu", "Temel", "SOC 2 Type II", "🔴 P0"],
            ["SLA Monitoring", "Yok", "%99.9 uptime garantisi", "🔴 P0"],
            ["Backup/Recovery", "Manuel", "Otomatik, point-in-time", "🔴 P0"]
        ]
    )
    
    add_heading_with_style(doc, "7.2 Önemli Eksiklikler (Should-Have)", 2)
    add_table_with_style(doc,
        ["Gap", "Mevcut", "Gerekli", "Öncelik"],
        [
            ["Observability", "%30", "Full stack (Prometheus + Grafana)", "🟡 P1"],
            ["API Rate Limiting", "Temel", "Per-tenant, tiered", "🟡 P1"],
            ["Audit Logging", "Kısmi", "Tam compliance logging", "🟡 P1"],
            ["Admin Dashboard", "Temel", "Multi-tenant yönetim", "🟡 P1"],
            ["SSO Integration", "Yok", "SAML/OIDC", "🟡 P1"],
            ["Billing Integration", "Yok", "Usage-based billing", "🟡 P1"]
        ]
    )
    
    doc.add_page_break()
    
    # =========================================================================
    # 8. EVRİM YOL HARİTASI
    # =========================================================================
    add_heading_with_style(doc, "8. Evrim Yol Haritası", 1)
    
    add_heading_with_style(doc, "Faz 0: Stabilizasyon (2 Hafta)", 2)
    doc.add_paragraph("Hedef: Mevcut sistemi enterprise-ready hale getirmek için temel düzeltmeler")
    add_table_with_style(doc,
        ["İş", "Detay", "Süre"],
        [
            ["Technical Debt Temizliği", "Test coverage %80+, lint fix", "3 gün"],
            ["SQLite → PostgreSQL Migrasyonu", "Alembic migrations", "3 gün"],
            ["Redis Cluster Kurulumu", "High-availability", "2 gün"],
            ["Backup Stratejisi", "Otomatik daily backup", "2 gün"],
            ["CI/CD Pipeline", "GitHub Actions", "2 gün"]
        ]
    )
    
    add_heading_with_style(doc, "Faz 1: Multi-Tenant Foundation (4 Hafta)", 2)
    doc.add_paragraph("Hedef: 10 müşteri için izole ortamlar")
    add_table_with_style(doc,
        ["İş", "Detay", "Süre"],
        [
            ["Tenant Model Design", "tenant_id propagation", "1 hafta"],
            ["Database Schema Update", "Row-level security", "1 hafta"],
            ["API Authentication", "JWT + tenant claims", "3 gün"],
            ["Namespace Isolation", "Redis + ChromaDB", "4 gün"],
            ["Kubernetes Deployment", "Helm charts", "1 hafta"]
        ]
    )
    
    add_heading_with_style(doc, "Faz 2: Enterprise Features (6 Hafta)", 2)
    add_table_with_style(doc,
        ["Hafta", "İş Paketi"],
        [
            ["1-2", "Full Observability Stack (Prometheus, Grafana, Jaeger)"],
            ["2-3", "Audit Logging & Compliance (GDPR, KVKK)"],
            ["3-4", "SSO Integration (SAML 2.0, OIDC)"],
            ["4-5", "Admin Dashboard v2 (Multi-tenant management)"],
            ["5-6", "Billing & Usage Metering"]
        ]
    )
    
    add_heading_with_style(doc, "Faz 3: AI Quality Enhancement (4 Hafta)", 2)
    doc.add_paragraph("Hedef: ChatGPT/Claude/Gemini seviyesine yaklaşmak")
    add_table_with_style(doc,
        ["İş", "Detay", "Beklenen İyileşme"],
        [
            ["Model Catalog Expansion", "GPT-4o, Claude 3.5, Gemini 2.0", "Kalite %30↑"],
            ["Adaptive Model Selection", "Task-based routing", "Latency %20↓"],
            ["Advanced RAG", "Agentic RAG, multi-hop reasoning", "Accuracy %25↑"],
            ["Memory Enhancement", "Graph-based relationships", "Context %40↑"],
            ["Proactive Assistant", "Suggestion engine", "Engagement %35↑"]
        ]
    )
    
    add_heading_with_style(doc, "Faz 4: Premium Features (8 Hafta)", 2)
    add_table_with_style(doc,
        ["İş", "Detay"],
        [
            ["Voice AI", "Whisper + TTS entegrasyonu"],
            ["Vision AI", "GPT-4V / Claude Vision"],
            ["Code Assistant", "GitHub/GitLab entegrasyonu, code review"],
            ["Document Intelligence", "Advanced PDF analysis, table extraction"],
            ["Workflow Builder", "Visual automation designer"]
        ]
    )
    
    doc.add_page_break()
    
    # =========================================================================
    # 9. TEKNİK BORÇ VE RİSKLER
    # =========================================================================
    add_heading_with_style(doc, "9. Teknik Borç ve Riskler", 1)
    
    add_heading_with_style(doc, "Yüksek Öncelikli Teknik Borç", 2)
    add_table_with_style(doc,
        ["Borç", "Etki", "Çözüm Süresi"],
        [
            ["SQLite production'da", "Ölçeklenebilirlik", "3 gün"],
            ["Monolithic deployment", "Single point of failure", "1 hafta"],
            ["Test coverage <%50", "Regression riski", "1 hafta"],
            ["Hardcoded API keys", "Güvenlik", "1 gün"],
            ["Missing rate limiting", "DDoS riski", "2 gün"]
        ]
    )
    
    add_heading_with_style(doc, "Riskler", 2)
    add_table_with_style(doc,
        ["Risk", "Olasılık", "Etki", "Mitigasyon"],
        [
            ["Groq API rate limits", "Yüksek", "Kritik", "4+ API key rotation"],
            ["Model deprecation", "Orta", "Yüksek", "Capability-based routing"],
            ["Data breach", "Düşük", "Kritik", "Encryption at rest/transit"],
            ["ChromaDB scaling", "Orta", "Orta", "Pinecone/Qdrant geçişi"]
        ]
    )
    
    doc.add_page_break()
    
    # =========================================================================
    # 10. ÖNERİLER VE SONUÇ
    # =========================================================================
    add_heading_with_style(doc, "10. Öneriler ve Sonuç", 1)
    
    add_heading_with_style(doc, "Acil Aksiyonlar (İlk 2 Hafta)", 2)
    actions = [
        "PostgreSQL Migrasyonu - SQLite enterprise ölçekte yetersiz",
        "CI/CD Pipeline - Otomatik test ve deployment",
        "Secrets Management - HashiCorp Vault veya AWS Secrets Manager",
        "Basic Monitoring - En azından uptime ve error rate"
    ]
    for i, action in enumerate(actions, 1):
        doc.add_paragraph(f"{i}. {action}")
    
    add_heading_with_style(doc, "Kısa Vadeli (1-3 Ay)", 2)
    short_term = [
        "Multi-tenant Mimari - 10 müşteri izolasyonu",
        "Kubernetes Deployment - Ölçeklenebilirlik",
        "Full Observability - Prometheus + Grafana + Jaeger",
        "SSO Entegrasyonu - Enterprise müşteriler için şart"
    ]
    for i, item in enumerate(short_term, 1):
        doc.add_paragraph(f"{i}. {item}")
    
    add_heading_with_style(doc, "Sonuç", 2)
    conclusion = doc.add_paragraph()
    conclusion.add_run("Mami AI v4.2").bold = True
    conclusion.add_run(
        ", güçlü bir teknik temele sahip ve birçok gelişmiş özelliği zaten barındırmaktadır. "
        "Orchestrator mimarisi, hafıza sistemleri ve RAG altyapısı enterprise-grade kalitededir."
    )
    
    doc.add_paragraph()
    
    warning = doc.add_paragraph()
    warning.add_run("10 premium müşteriye hizmet için:\n").bold = True
    warning.add_run("• Multi-tenant izolasyon zorunludur\n")
    warning.add_run("• PostgreSQL geçişi kritiktir\n")
    warning.add_run("• Güvenlik sertifikasyonu (en az SOC 2) gereklidir\n")
    warning.add_run("• SLA garantisi için monitoring altyapısı şarttır")
    
    doc.add_paragraph()
    
    estimate = doc.add_paragraph()
    estimate.add_run("Tahmini Geliştirme Süresi: ").bold = True
    estimate.add_run("4-6 ay (4 kişilik ekip ile)\n\n")
    estimate.add_run("Tahmini Maliyet Kalemleri:\n").bold = True
    estimate.add_run("• Cloud Infrastructure: ~$2,000-5,000/ay\n")
    estimate.add_run("• LLM API Costs: ~$1,000-3,000/ay (10 müşteri)\n")
    estimate.add_run("• Monitoring Tools: ~$500/ay\n")
    estimate.add_run("• Geliştirme Ekibi: 4 Senior Engineer")
    
    doc.add_page_break()
    
    # =========================================================================
    # EKLER
    # =========================================================================
    add_heading_with_style(doc, "Ekler", 1)
    
    add_heading_with_style(doc, "Ek A: Dosya İstatistikleri", 2)
    add_table_with_style(doc,
        ["Kategori", "Dosya Sayısı", "Toplam Satır"],
        [
            ["Backend Python", "150+", "~25,000"],
            ["Frontend TSX/TS", "100+", "~15,000"],
            ["Tests", "42", "~3,000"],
            ["Docs", "30+", "~5,000"],
            ["TOPLAM", "300+", "~50,000"]
        ]
    )
    
    add_heading_with_style(doc, "Ek B: Model Catalog (Mevcut)", 2)
    add_table_with_style(doc,
        ["Model", "Provider", "Use Case"],
        [
            ["llama-3.3-70b-versatile", "Groq", "Ana yanıt üretimi"],
            ["llama-3.1-8b-instant", "Groq", "Hızlı işlemler, routing"],
            ["josiefied-qwen3-8b", "Ollama", "Sansürsüz içerik"],
            ["Flux", "Forge", "Görsel üretim (SFW)"],
            ["FluxedUp NSFW", "Forge", "Görsel üretim (NSFW)"]
        ]
    )
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("─" * 50 + "\n").font.color.rgb = RGBColor(150, 150, 150)
    footer.add_run(
        "Bu belge, Mami AI v4.2 projesinin kapsamlı teknik değerlendirmesini içermektedir.\n"
        "Enterprise deployment öncesi tüm önerilerin dikkatle değerlendirilmesi tavsiye edilir."
    ).font.color.rgb = RGBColor(100, 100, 100)
    
    # Save
    output_path = Path("d:/ai/mami_ai_v4/docs/Mami_AI_Enterprise_Degerlendirme_Raporu.docx")
    doc.save(output_path)
    print(f"✅ Word belgesi oluşturuldu: {output_path}")
    return output_path

if __name__ == "__main__":
    create_enterprise_report()
